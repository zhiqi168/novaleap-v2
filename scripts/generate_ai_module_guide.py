# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("D:/Desktop/novaleap")
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

OUT_CN = DOCS / "AI模块完整流程与逐句讲解-修复版.docx"
OUT_ASCII = DOCS / "ai-module-guide-fixed.docx"

FILES = [
    ("后端入口", "nova-backend/src/main/java/com/novaleap/api/controller/AiController.java"),
    ("后端接口定义", "nova-backend/src/main/java/com/novaleap/api/service/AiService.java"),
    ("后端业务编排核心", "nova-backend/src/main/java/com/novaleap/api/service/impl/AiServiceImpl.java"),
    ("模型调用网关", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiModelGateway.java"),
    ("Prompt 生成工厂", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiPromptFactory.java"),
    ("AI 对话历史与会话", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiCoachSessionSupport.java"),
    ("AI 外部联网上下文", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiExternalContextService.java"),
    ("笔记摘要与审核流程", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiNoteWorkflowSupport.java"),
    ("内容解析与规则兜底", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiContentSupport.java"),
    ("身份与输入辅助", "nova-backend/src/main/java/com/novaleap/api/module/ai/support/AiIdentitySupport.java"),
    ("AI 调用审计", "nova-backend/src/main/java/com/novaleap/api/module/ai/audit/AiCallAuditService.java"),
    ("AI 网关配置对象", "nova-backend/src/main/java/com/novaleap/api/module/ai/config/AiGatewayProperties.java"),
    ("AI 限流接口", "nova-backend/src/main/java/com/novaleap/api/service/AiLimitService.java"),
    ("AI 限流实现", "nova-backend/src/main/java/com/novaleap/api/service/impl/AiLimitServiceImpl.java"),
    ("AI 配额策略", "nova-backend/src/main/java/com/novaleap/api/module/quota/support/AiQuotaPolicy.java"),
    ("AI Token 用量记录", "nova-backend/src/main/java/com/novaleap/api/module/quota/support/AiQuotaUsageSupport.java"),
    ("AI 配额配置对象", "nova-backend/src/main/java/com/novaleap/api/module/quota/config/AiQuotaProperties.java"),
    ("AI 视图组装", "nova-backend/src/main/java/com/novaleap/api/module/ai/assembler/AiViewAssembler.java"),
    ("请求 DTO：教练聊天", "nova-backend/src/main/java/com/novaleap/api/module/ai/dto/AiCoachChatRequest.java"),
    ("请求 DTO：笔记摘要", "nova-backend/src/main/java/com/novaleap/api/module/ai/dto/AiNoteSummaryRequest.java"),
    ("请求 DTO：简历分析", "nova-backend/src/main/java/com/novaleap/api/module/ai/dto/AiResumeAnalyzeRequest.java"),
    ("返回 VO：历史记录", "nova-backend/src/main/java/com/novaleap/api/module/ai/vo/AiCoachHistoryItemVO.java"),
    ("返回 VO：会话", "nova-backend/src/main/java/com/novaleap/api/module/ai/vo/AiCoachSessionVO.java"),
    ("前端 AI 聊天页面", "nova-frontend/src/views/Coach.vue"),
    ("前端简历分析页面", "nova-frontend/src/views/Resume.vue"),
    ("前端题目 AI 抽屉", "nova-frontend/src/components/common/AiDrawer.vue"),
    ("前端摘要展示块", "nova-frontend/src/components/common/AiSummaryBlock.vue"),
    ("前端 SSE 流式接收", "nova-frontend/src/composables/useSSE.js"),
]

FILE_ROLE = {
    "AiController.java": "所有 /api/ai 开头请求的入口。它负责接收前端参数、识别当前用户，然后把事情交给 AiService。",
    "AiService.java": "AI 能力的接口清单。你可以把它当成菜单：题目解释、教练聊天、简历分析、笔记摘要、每日一句、笔记审核。",
    "AiServiceImpl.java": "AI 业务编排核心。它负责输入校验、限流、身份问答、读取历史、拼 prompt、调用模型、失败兜底。",
    "AiModelGateway.java": "真正调用大模型的网关。这里处理 Spring AI ChatClient、SSE 流式输出、超时、重试、熔断、备用模型和 token 统计。",
    "AiPromptFactory.java": "集中生成 prompt。想改 AI 的回答风格、结构、身份话术，优先改这里。",
    "AiCoachSessionSupport.java": "把聊天记录和当前会话存进 Redis，多轮对话上下文主要靠它。",
    "AiExternalContextService.java": "给 AI 补充外部信息，例如搜索摘要和天气信息，再塞进 prompt。",
    "AiNoteWorkflowSupport.java": "笔记摘要和笔记审核的流程层。它先做规则兜底，再尝试调用 AI。",
    "AiContentSupport.java": "处理摘要解析、敏感词、审核 JSON 解析，以及 AI 不可用时的规则兜底。",
    "AiIdentitySupport.java": "处理用户身份、输入清洗，并判断用户是不是在问 AI 身份或创作者。",
    "AiCallAuditService.java": "记录 AI 调用成功/失败和 token 用量，方便排查模型稳定性。",
    "AiGatewayProperties.java": "把 nova.ai.gateway 配置映射成 Java 对象，例如超时、重试次数、熔断窗口。",
    "AiLimitService.java": "AI 限流服务接口，定义模块类型和限流检查结果。",
    "AiLimitServiceImpl.java": "AI 限流实现。它用 Redis 记录调用次数、冷却时间、每日 token 使用等。",
    "AiQuotaPolicy.java": "不同角色、不同 AI 模块每天能用多少次，以及 token 用量到多少进入降级。",
    "AiQuotaUsageSupport.java": "记录当天 AI token 总消耗，配合降级和保护策略。",
    "AiQuotaProperties.java": "把 nova.ai.quota 配置映射成 Java 对象，例如每日 token 上限、预警比例。",
    "AiViewAssembler.java": "把后端 Map 或实体转换成前端更好用的 VO。",
    "AiCoachChatRequest.java": "前端发起 AI 教练聊天时传来的 JSON 请求结构。",
    "AiNoteSummaryRequest.java": "前端请求笔记摘要时传来的 JSON 请求结构。",
    "AiResumeAnalyzeRequest.java": "前端请求简历分析时传来的 JSON 请求结构。",
    "AiCoachHistoryItemVO.java": "返回给前端的一条聊天历史记录。",
    "AiCoachSessionVO.java": "创建新会话后返回给前端的会话 id。",
    "Coach.vue": "AI 聊天页面。负责输入、图片压缩、发送请求、展示流式回复和恢复历史。",
    "Resume.vue": "简历分析页面。把简历文本和目标岗位发给后端，展示流式分析结果。",
    "AiDrawer.vue": "题目页面侧边 AI 抽屉，用来展示题目解释的流式结果。",
    "AiSummaryBlock.vue": "笔记摘要展示组件。",
    "useSSE.js": "前端接收 text/event-stream 的通用工具。所有流式 AI 输出都靠它拼接。",
}


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def set_font(run, name="Microsoft YaHei", size=10, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc: Document, text: str, size=10.5, bold=False, color=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_heading(level=level)
    p.text = ""
    r = p.add_run(text)
    set_font(r, size={0: 22, 1: 16, 2: 13, 3: 11}.get(level, 11), bold=True, color=(31, 78, 121))
    return p


def add_bullet(doc: Document, text: str):
    add_para(doc, text, size=10, style="List Bullet")


def add_number(doc: Document, text: str):
    add_para(doc, text, size=10, style="List Number")


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def explain_line(line: str, filename: str) -> str:
    s = line.strip()
    if not s:
        return "空行：把代码分隔开，让结构更清楚。"

    if filename.endswith(".vue"):
        if s.startswith("<template"):
            return "Vue 模板开始，下面写页面上真正渲染出来的结构。"
        if s.startswith("</template"):
            return "Vue 模板结束。"
        if s.startswith("<script"):
            return "Vue 脚本逻辑开始，变量和函数都在这里定义。"
        if s.startswith("</script"):
            return "脚本区域结束。"
        if s.startswith("<style"):
            return "组件样式开始，scoped 表示样式主要影响当前组件。"
        if s.startswith("</style"):
            return "样式区域结束。"
        if s.startswith("<") and not s.startswith("</"):
            tag = re.match(r"<([\w-]+)", s)
            tag_name = tag.group(1) if tag else "标签"
            if "v-if" in s:
                return f"渲染一个 {tag_name}，并用 v-if 控制是否显示。"
            if "v-for" in s:
                return f"渲染一个 {tag_name}，并用 v-for 根据数组循环生成内容。"
            if "@click" in s:
                return f"渲染一个 {tag_name}，并绑定点击事件。"
            if "v-model" in s:
                return f"渲染一个 {tag_name} 输入控件，并和变量双向绑定。"
            if ":" in s:
                return f"渲染一个 {tag_name}，其中冒号开头的属性会随变量动态变化。"
            return f"渲染一个 {tag_name}，属于当前页面结构。"
        if s.startswith("</"):
            return "结束前面打开的 Vue/HTML 标签。"

    if filename.endswith(".js") or filename.endswith(".vue"):
        if s.startswith("import "):
            return "导入前端组件、工具函数或依赖。"
        if s.startswith("const ") or s.startswith("let "):
            if "ref(" in s:
                return "定义 Vue 响应式变量；值变化后页面会自动更新。"
            if "computed(" in s:
                return "定义计算属性，会根据依赖变量自动算出结果。"
            if "useSSE" in s:
                return "调用 SSE 工具，拿到流式内容、加载状态、错误状态和启动/中断方法。"
            return "定义前端变量或常量，用来保存数据、函数或配置。"
        if s.startswith("watch("):
            return "监听响应式变量变化；AI 流式内容变化时同步更新页面。"
        if s.startswith("onMounted("):
            return "组件加载完成后执行初始化逻辑。"
        if "fetch(" in s:
            return "前端发起 HTTP 请求到后端。"
        if "reader.read()" in s:
            return "从后端响应流里读取下一段数据，这是接收 AI 流式回复的关键。"
        if "JSON.parse" in s:
            return "把后端发来的 JSON 字符串解析成对象。"
        if ".value" in s and "=" in s:
            return "修改 Vue ref 的值，页面会响应式更新。"

    if s.startswith("package "):
        return "声明这个 Java 文件属于哪个包，方便项目按模块组织代码。"
    if s.startswith("import "):
        return "引入外部类或项目内其他类，后面才能直接使用。"
    if s.startswith("@RestController"):
        return "告诉 Spring：这是 REST 接口控制器，负责接收 HTTP 请求。"
    if s.startswith("@RequestMapping"):
        return "设置接口路径前缀，这里的 AI 接口统一挂在 /api/ai 下。"
    if s.startswith("@GetMapping"):
        return "声明一个 GET 接口，通常用于读取数据或触发只读型流式结果。"
    if s.startswith("@PostMapping"):
        return "声明一个 POST 接口，前端会把 JSON 请求体发到这里。"
    if s.startswith("@DeleteMapping"):
        return "声明一个 DELETE 接口，用来清空或删除数据。"
    if s.startswith("@Service"):
        return "把这个类交给 Spring 管理，其他类可以自动注入它。"
    if s.startswith("@Component"):
        return "声明通用 Spring 组件，会被 Spring 自动创建对象。"
    if s.startswith("@ConfigurationProperties"):
        return "把 application.yml 中某个前缀下的配置绑定到这个类。"
    if s.startswith("@Validated"):
        return "开启配置或参数校验。"
    if s.startswith("@Value"):
        return "从配置文件或环境变量读取值，例如 API Key、模型名。"
    if s.startswith("@Override"):
        return "说明下面的方法是在实现接口或重写父类方法。"
    if s.startswith("@Min") or s.startswith("@Max") or s.startswith("@Not") or s.startswith("@Decimal"):
        return "校验注解，用来限制字段不能为空或数值范围必须合理。"
    if s.startswith("@Slf4j"):
        return "Lombok 自动提供 log 日志对象。"
    if re.match(r"public\s+class\s+", s):
        return "定义一个 Java 类，里面放字段、构造函数和方法。"
    if re.match(r"public\s+interface\s+", s):
        return "定义接口，只规定能力，不写具体实现。"
    if " record " in f" {s} " or s.startswith("public record"):
        return "定义 record 数据对象，适合保存简单字段并返回给前端。"
    if s == "{":
        return "代码块开始，下面内容属于上一行的类、方法、if 或循环。"
    if s == "}":
        return "代码块结束，上一层逻辑在这里收尾。"
    if s.startswith("private static final"):
        return "定义常量，例如默认值、错误提示、Redis key 前缀或固定文案。"
    if s.startswith("private final"):
        return "定义不可重新赋值的依赖字段，通常通过构造函数注入。"
    if re.match(r"(public|private|protected).*\)\s*\{", s):
        return "定义一个方法或构造函数。括号里是参数，大括号里是具体逻辑。"
    if s.startswith("if"):
        if "hasAiCapability" in s:
            return "判断当前 AI 能力是否可用，主要看 API Key 是否配置。"
        if "limit.isAllowed" in s:
            return "判断限流是否通过；不通过就返回提示，避免继续调用模型。"
        if "isBlank" in s or "isEmpty" in s or "null" in s:
            return "做空值保护，避免没有输入或没有数据时继续往下跑。"
        if "isCircuitOpen" in s:
            return "判断模型是否熔断；熔断表示最近失败太多，先别打这个模型。"
        return "条件判断：只有满足条件才执行里面的代码。"
    if s.startswith("else"):
        return "前面的条件不满足时走这里，通常是兜底分支。"
    if s.startswith("for") or s.startswith("while"):
        return "循环处理多条数据，例如历史记录、流式片段或列表项。"
    if s.startswith("try"):
        return "开始异常保护；里面出错会进入 catch，不让程序直接崩。"
    if s.startswith("catch"):
        return "捕获错误，记录日志或返回兜底结果。"
    if s.startswith("finally"):
        return "无论成功失败都会执行，常用于清理状态。"
    if s.startswith("return "):
        if "Result.success" in s:
            return "返回统一成功响应给前端。"
        if "streamStaticAnswer" in s:
            return "返回固定文本的伪流式响应，用于错误、限流或 AI 不可用提示。"
        if "streamModelAnswer" in s:
            return "返回真正模型生成的 SSE 流式响应。"
        return "结束当前方法，把结果交给调用方。"
    if "aiLimitService.checkLimit" in s:
        return "调用限流服务，检查这个用户当前模块还能不能用 AI。"
    if "aiPromptFactory" in s:
        return "调用 Prompt 工厂生成提示词，告诉模型该怎么回答。"
    if "aiModelGateway.streamModelAnswer" in s:
        return "调用模型网关进行流式生成，结果会一段段推给前端。"
    if "aiModelGateway.callModel" in s:
        return "调用模型网关进行普通非流式生成，等完整结果回来。"
    if "saveCoachMessage" in s:
        return "把聊天消息写入 Redis 历史，后面可以作为上下文。"
    if "getCoachHistory" in s:
        return "读取最近聊天历史，用于恢复页面或拼多轮上下文。"
    if "SseEmitter" in s:
        return "Spring 的 SSE 对象，用来持续向前端发送事件流。"
    if "ChatClient" in s or "chatClient.prompt" in s:
        return "Spring AI 的模型调用入口，用来构造聊天请求。"
    if ".system(" in s:
        return "设置 system prompt，也就是模型必须遵守的角色和规则。"
    if ".user(" in s:
        return "设置 user prompt，也就是用户问题和业务上下文。"
    if ".options(" in s or "OpenAiChatOptions" in s:
        return "设置模型调用参数，例如模型名称、最大输出 token。"
    if ".stream()" in s:
        return "开启流式模型调用，模型边生成边返回。"
    if ".call()" in s:
        return "普通模型调用，等待完整回答后一次性返回。"
    if ".subscribe(" in s:
        return "订阅流式结果：收到片段、出错、完成时分别处理。"
    if "redisTemplate" in s or "StringRedisTemplate" in s:
        return "操作 Redis，用于历史、限流、token 用量或熔断状态。"
    if "objectMapper" in s:
        return "处理 JSON，把字符串和对象互相转换。"
    if "CompletableFuture" in s or "ExecutorService" in s:
        return "使用后台线程执行任务，避免阻塞主请求线程。"
    if "Duration." in s or "expire(" in s:
        return "设置 Redis key 的过期时间，防止历史记录和计数永久堆积。"
    if "log." in s:
        return "写日志，方便排查问题。"
    if s.startswith("//") or s.startswith("*") or s.startswith("/*"):
        return "注释，用自然语言说明附近代码的目的。"
    if "=" in s:
        return "赋值：把右边的结果保存到左边变量或字段。"
    return "具体业务代码：结合上下文参与数据准备、判断、调用或结果收尾。"


def add_code_table(doc: Document, path: Path) -> int:
    lines = read_text(path).splitlines()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    repeat_header(header)
    for idx, title in enumerate(("行号", "代码", "通俗解释")):
        cell = header.cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(title)
        set_font(run, size=9, bold=True, color=(255, 255, 255))
        shade(cell, "2F5597")

    for no, line in enumerate(lines, start=1):
        row = table.add_row()
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        run = row.cells[0].paragraphs[0].add_run(str(no))
        set_font(run, size=8)

        run = row.cells[1].paragraphs[0].add_run(line if line.strip() else " ")
        set_font(run, name="Consolas", size=7.5)

        run = row.cells[2].paragraphs[0].add_run(explain_line(line, path.name))
        set_font(run, size=8.5)

    return len(lines)


def build_doc() -> int:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("NovaLeap AI 模块完整流程与逐句讲解")
    set_font(run, size=22, bold=True, color=(31, 78, 121))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M"))
    set_font(run, size=10, color=(100, 100, 100))

    add_para(doc, "这份文档按项目真实代码讲解 AI 模块：先看完整流程，再看调用链，最后逐文件逐句解释每行代码的意思。", size=11)

    add_heading(doc, "一、完整流程人话版")
    for item in [
        "用户在 Coach.vue 输入消息或上传图片，前端整理成 JSON。",
        "Coach.vue 调用 useSSE.js 的 startStream，向 /api/ai/coach/chat 发 POST 请求。",
        "AiController.coachChat 接收请求，解析当前用户或 IP。",
        "AiServiceImpl.coachChat 校验输入、检查 API Key、检查限流。",
        "如果问身份或创作者，代码走固定回复，避免模型乱说。",
        "正常聊天时，从 Redis 取历史记录，可能再取搜索/天气等外部上下文。",
        "AiPromptFactory 把主题、模式、历史、问题、外部上下文拼成 prompt。",
        "AiModelGateway 用 Spring AI ChatClient 调模型，并通过 SseEmitter 一段段返回。",
        "useSSE.js 读取 text/event-stream，把 delta 内容拼起来。",
        "Coach.vue 监听 content，更新最后一条 AI 消息，于是页面出现打字效果。",
    ]:
        add_number(doc, item)

    add_heading(doc, "二、核心调用链")
    for item in [
        "聊天：Coach.vue -> useSSE.js -> AiController.coachChat -> AiServiceImpl.coachChat -> AiPromptFactory -> AiModelGateway -> ChatClient -> SseEmitter -> useSSE.js -> Coach.vue",
        "题目解释：AiDrawer.vue -> useSSE.js -> AiController.explainQuestion -> AiServiceImpl.explainQuestion -> AiPromptFactory.buildQuestionExplainPrompt -> AiModelGateway.streamModelAnswer",
        "简历分析：Resume.vue -> useSSE.js -> AiController.analyzeResume -> AiServiceImpl.analyzeResume -> AiPromptFactory.buildResumePrompt -> AiModelGateway.streamModelAnswer",
        "笔记摘要：前端摘要组件 -> AiController.summarizeNote -> AiServiceImpl.summarizeNote -> AiNoteWorkflowSupport -> AiModelGateway.callModelRaw",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、先懂这几个词")
    for item in [
        "SSE：后端一段段推送文本给前端，所以 AI 像打字。",
        "Prompt：给模型的任务说明，你的项目把问题、历史和上下文都拼进去。",
        "System Prompt：模型必须遵守的角色规则，例如 NovaLeap 身份。",
        "Redis：保存聊天历史、限流计数、token 用量、熔断状态。",
        "限流：控制每个用户每天能用多少次 AI。",
        "熔断：模型连续失败时短时间避开它，保护系统稳定。",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "四、逐文件逐句讲解")

    total_lines = 0
    for category, rel_path in FILES:
        path = ROOT / rel_path
        if not path.exists():
            continue
        add_heading(doc, f"{category}：{path.name}", level=2)
        add_para(doc, f"路径：{path}", size=9, color=(90, 90, 90))
        add_para(doc, "作用：" + FILE_ROLE.get(path.name, "AI 模块的辅助文件，参与请求、响应、配置或展示。"), size=10.5)
        count = add_code_table(doc, path)
        total_lines += count
        add_para(doc, f"本文件讲解 {count} 行。", size=9, color=(90, 90, 90))
        doc.add_page_break()

    add_heading(doc, "五、改代码时优先看哪里")
    for item in [
        "改 AI 回答风格：AiPromptFactory.java",
        "改聊天主流程：AiServiceImpl.java",
        "改模型调用、超时、重试、熔断：AiModelGateway.java",
        "改限流次数：AiQuotaPolicy.java / AiLimitServiceImpl.java",
        "改聊天页面体验：Coach.vue",
        "改流式接收和错误提示：useSSE.js",
    ]:
        add_bullet(doc, item)

    add_para(doc, f"本 Word 共逐行覆盖约 {total_lines} 行代码。", bold=True)
    doc.save(OUT_CN)
    shutil.copyfile(OUT_CN, OUT_ASCII)
    return total_lines


if __name__ == "__main__":
    lines = build_doc()
    print(str(OUT_CN))
    print(str(OUT_ASCII))
    print(lines)
